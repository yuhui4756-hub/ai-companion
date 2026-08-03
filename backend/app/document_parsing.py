from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from typing import Any

from .schemas import KnowledgeSourceType

MAX_FILE_BYTES = 5 * 1024 * 1024
MAX_PARSED_CHARS = 190_000
MAX_PDF_PAGES = 80
MAX_DOCX_BLOCKS = 2_000

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}


class DocumentParseError(ValueError):
    def __init__(self, message: str, *, status_code: int = 400) -> None:
        super().__init__(message)
        self.status_code = status_code


@dataclass(frozen=True)
class ParsedDocument:
    title: str
    source_type: KnowledgeSourceType
    content: str
    metadata: dict[str, Any]


def clean_file_name(filename: str | None) -> str:
    cleaned = (filename or "").replace("\\", "/").split("/")[-1].strip()
    return cleaned or "未命名资料"


def file_extension(filename: str) -> str:
    lowered = filename.lower()
    if "." not in lowered:
        return ""
    return f".{lowered.rsplit('.', 1)[-1]}"


def title_from_filename(filename: str) -> str:
    base = filename.rsplit(".", 1)[0].strip() if "." in filename else filename.strip()
    return base[:120] or "未命名资料"


def ensure_file_size(data: bytes) -> None:
    if not data:
        raise DocumentParseError("文件内容为空，请选择包含文字的资料。")
    if len(data) > MAX_FILE_BYTES:
        raise DocumentParseError("文件过大，本轮最多支持 5MB 的文本层资料。", status_code=413)


def normalize_text(value: str) -> str:
    text = value.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    normalized = "\n".join(lines)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized).strip()
    return normalized


def ensure_parsed_text_budget(text: str) -> None:
    if not text.strip():
        raise DocumentParseError("没有解析到可导入的文本内容。")
    if len(text) > MAX_PARSED_CHARS:
        raise DocumentParseError("解析后的文本过长，本轮最多支持约 19 万字符，请拆分后再导入。", status_code=413)


def decode_utf8_text(data: bytes) -> str:
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise DocumentParseError("文本文件必须是 UTF-8 编码；请转换编码后再导入。") from error


def markdown_escape_cell(value: str) -> str:
    return normalize_text(value).replace("|", "\\|")


def parse_text_document(filename: str, data: bytes, source_type: KnowledgeSourceType) -> ParsedDocument:
    content = normalize_text(decode_utf8_text(data))
    ensure_parsed_text_budget(content)
    extension = file_extension(filename)
    return ParsedDocument(
        title=title_from_filename(filename),
        source_type=source_type,
        content=content,
        metadata={
            "sourceFormat": "markdown" if extension in {".md", ".markdown"} else "plain_text_file",
            "fileName": filename,
            "fileSizeBytes": len(data),
            "originalExtension": extension,
        },
    )


def parse_pdf_document(filename: str, data: bytes) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise DocumentParseError("当前环境缺少 PDF 解析依赖，请先安装后端依赖。", status_code=500) from error

    try:
        reader = PdfReader(BytesIO(data))
    except Exception as error:
        raise DocumentParseError("PDF 解析失败，请确认文件未损坏且包含文本层。") from error

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as error:
            raise DocumentParseError("暂不支持加密 PDF，请先解除密码后再导入。") from error
        if reader.is_encrypted:
            raise DocumentParseError("暂不支持加密 PDF，请先解除密码后再导入。")

    page_count = len(reader.pages)
    if page_count > MAX_PDF_PAGES:
        raise DocumentParseError(f"PDF 页数超过本轮上限 {MAX_PDF_PAGES} 页，请拆分后再导入。", status_code=413)

    title = title_from_filename(filename)
    parts = [f"# {title}"]
    extracted_pages = 0
    for index, page in enumerate(reader.pages, start=1):
        try:
            page_text = normalize_text(page.extract_text() or "")
        except Exception:
            page_text = ""
        if not page_text:
            continue
        extracted_pages += 1
        parts.extend([f"## 第 {index} 页", page_text])

    if extracted_pages == 0:
        raise DocumentParseError("当前暂不支持自动 OCR，可复制文本或手动补充说明。")

    content = "\n\n".join(parts)
    ensure_parsed_text_budget(content)
    return ParsedDocument(
        title=title,
        source_type=KnowledgeSourceType.pdf_text,
        content=content,
        metadata={
            "sourceFormat": "pdf_text",
            "fileName": filename,
            "fileSizeBytes": len(data),
            "originalExtension": ".pdf",
            "pageCount": page_count,
            "textPageCount": extracted_pages,
        },
    )


def docx_paragraph_to_markdown(paragraph: Any) -> str:
    text = normalize_text(paragraph.text or "")
    if not text:
        return ""

    style_name = (getattr(paragraph.style, "name", "") or "").strip().lower()
    heading_match = re.search(r"(?:heading|标题)\s*(\d+)", style_name)
    if heading_match:
        level = min(max(int(heading_match.group(1)), 1), 6)
        return f"{'#' * level} {text}"
    if "list" in style_name or "列表" in style_name:
        return f"- {text}"
    return text


def docx_table_to_markdown_rows(table: Any, table_index: int) -> list[str]:
    rows = table.rows
    if not rows:
        return []

    first_cells = [normalize_text(cell.text or "") for cell in rows[0].cells]
    has_header = any(first_cells)
    headers = [value or f"列{index + 1}" for index, value in enumerate(first_cells)] if has_header else []
    blocks = [f"## 表格 {table_index}"]
    data_rows = rows[1:] if has_header and len(rows) > 1 else rows

    for row_index, row in enumerate(data_rows, start=1):
        values = [normalize_text(cell.text or "") for cell in row.cells]
        if not any(values):
            continue
        if not headers:
            headers = [f"列{index + 1}" for index in range(len(values))]
        row_lines = [f"### 表格 {table_index} 第 {row_index} 行"]
        for column_index, value in enumerate(values):
            header = headers[column_index] if column_index < len(headers) else f"列{column_index + 1}"
            if value:
                row_lines.append(f"{markdown_escape_cell(header)}：{markdown_escape_cell(value)}")
        blocks.append("\n".join(row_lines))
    return blocks


def iter_docx_blocks(document: Any) -> list[str]:
    try:
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as error:
        raise DocumentParseError("当前环境缺少 DOCX 解析依赖，请先安装后端依赖。", status_code=500) from error

    blocks: list[str] = []
    table_index = 0
    block_count = 0
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            markdown = docx_paragraph_to_markdown(paragraph)
            if markdown:
                blocks.append(markdown)
                block_count += 1
        elif child.tag.endswith("}tbl"):
            table_index += 1
            table = Table(child, document)
            table_blocks = docx_table_to_markdown_rows(table, table_index)
            blocks.extend(table_blocks)
            block_count += len(table_blocks)
        if block_count > MAX_DOCX_BLOCKS:
            raise DocumentParseError("DOCX 内容块过多，本轮请拆分后再导入。", status_code=413)
    return blocks


def parse_docx_document(filename: str, data: bytes) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as error:
        raise DocumentParseError("当前环境缺少 DOCX 解析依赖，请先安装后端依赖。", status_code=500) from error

    try:
        document = Document(BytesIO(data))
    except Exception as error:
        raise DocumentParseError("DOCX 解析失败，请确认文件未损坏且不是旧版 .doc 文件。") from error

    title = title_from_filename(filename)
    blocks = iter_docx_blocks(document)
    content = "\n".join([f"# {title}", *blocks]).strip()
    ensure_parsed_text_budget(content)
    return ParsedDocument(
        title=title,
        source_type=KnowledgeSourceType.docx,
        content=content,
        metadata={
            "sourceFormat": "docx",
            "fileName": filename,
            "fileSizeBytes": len(data),
            "originalExtension": ".docx",
            "blockCount": len(blocks),
        },
    )


def parse_document_bytes(filename: str | None, data: bytes) -> ParsedDocument:
    clean_name = clean_file_name(filename)
    extension = file_extension(clean_name)
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError("仅支持 .txt、.md、.pdf、.docx 文件。")

    ensure_file_size(data)
    if extension == ".txt":
        return parse_text_document(clean_name, data, KnowledgeSourceType.plain_text_file)
    if extension in {".md", ".markdown"}:
        return parse_text_document(clean_name, data, KnowledgeSourceType.markdown)
    if extension == ".pdf":
        return parse_pdf_document(clean_name, data)
    if extension == ".docx":
        return parse_docx_document(clean_name, data)
    raise DocumentParseError("不支持的文件类型。")
