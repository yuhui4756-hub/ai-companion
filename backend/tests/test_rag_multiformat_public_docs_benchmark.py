from __future__ import annotations

import json
import sqlite3
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from fastapi.testclient import TestClient

from backend.app.main import app
from backend.tests.test_document_parsing import make_text_pdf
from backend.tests.test_rag_realistic_benchmark import BenchmarkCase, evaluate_case


PUBLIC_MULTIFORMAT_SOURCE_URLS = {
    "React useEffect 多格式摘要": "https://react.dev/reference/react/useEffect",
    "Vite 环境变量多格式摘要": "https://vite.dev/guide/env-and-mode",
    "Electron 安全边界多格式摘要": "https://www.electronjs.org/docs/latest/tutorial/security",
    "FastAPI TestClient 多格式摘要": "https://fastapi.tiangolo.com/tutorial/testing/",
    "SQLite FTS5 BM25 多格式摘要": "https://www.sqlite.org/fts5.html",
    "Ollama Embeddings 多格式摘要": "https://docs.ollama.com/api",
    "DeepSeek API 多格式摘要": "https://api-docs.deepseek.com/",
    "MDN Fetch API 多格式摘要": "https://developer.mozilla.org/en-US/docs/Web/API/Fetch_API/Using_Fetch",
    "GitHub Actions Token 多格式摘要": "https://docs.github.com/en/actions/security-for-github-actions/security-guides/automatic-token-authentication",
    "Python venv 多格式摘要": "https://docs.python.org/3/library/venv.html",
}


@dataclass(frozen=True)
class FileFixture:
    title: str
    filename: str
    content: bytes
    media_type: str
    expected_source_type: str


def public_multiformat_runtime_config() -> dict:
    return {
        "providerName": "mock",
        "baseURL": "http://127.0.0.1:8765/mock",
        "model": "mock-embedding-multiformat-public",
        "dimensions": 64,
        "batchSize": 4,
        "timeoutMs": 3000,
        "enabled": True,
        "apiKey": "mock-local-key",
    }


def bytes_utf8(value: str) -> bytes:
    return value.strip().encode("utf-8")


def make_docx_bytes(title: str, sections: list[tuple[str, list[str]]], table_rows: list[dict[str, str]]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(title, level=1)
    for heading, lines in sections:
        document.add_heading(heading, level=2)
        for line in lines:
            document.add_paragraph(line)
    if table_rows:
        document.add_heading("表格字段", level=2)
        headers = list(table_rows[0].keys())
        table = document.add_table(rows=1, cols=len(headers))
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row_data in table_rows:
            row = table.add_row().cells
            for index, header in enumerate(headers):
                row[index].text = row_data[header]

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_public_multiformat_fixtures() -> list[FileFixture]:
    return [
        FileFixture(
            title="React useEffect 多格式摘要",
            filename="React useEffect 多格式摘要.md",
            media_type="text/markdown",
            expected_source_type="markdown",
            content=bytes_utf8(
                """
                # React useEffect 多格式摘要

                ## 生命周期
                术语：useEffect 是 React Hook，用来让组件和外部系统同步。
                setup 函数可以返回 cleanup 函数。
                依赖变化：React 会先用旧值运行 cleanup，再用新值运行 setup。
                卸载规则：组件从页面移除后，React 会最后运行 cleanup。
                依赖比较：dependencies 必须写成内联数组，React 使用 Object.is 比较每一项。
                边界：如果没有外部系统需要同步，通常不需要 Effect。
                """
            ),
        ),
        FileFixture(
            title="Vite 环境变量多格式摘要",
            filename="Vite 环境变量多格式摘要.txt",
            media_type="text/plain",
            expected_source_type="plain_text_file",
            content=bytes_utf8(
                """
                Vite 环境变量多格式摘要

                客户端变量：源码中通过 import.meta.env 读取环境变量。
                暴露规则：只有以 VITE_ 开头的变量会暴露给客户端代码。
                类型规则：环境变量值会以字符串形式暴露，业务代码需要自行转换布尔值或数字。
                重启规则：修改 .env 文件后需要重启开发服务器。
                内置字段：import.meta.env.MODE、DEV、PROD、SSR 和 BASE_URL 是 Vite 提供的内置字段。
                """
            ),
        ),
        FileFixture(
            title="Electron 安全边界多格式摘要",
            filename="Electron 安全边界多格式摘要.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            expected_source_type="docx",
            content=make_docx_bytes(
                "Electron 安全边界多格式摘要",
                [
                    (
                        "BrowserWindow 安全",
                        [
                            "nodeIntegration：加载远程内容的窗口不要启用 nodeIntegration。",
                            "contextIsolation：保持开启，让 preload 与网页运行在不同上下文。",
                            "sandbox：建议开启，减少渲染进程可用的特权能力。",
                            "webSecurity：不要关闭 webSecurity。",
                        ],
                    ),
                    (
                        "Preload 边界",
                        [
                            "只通过 contextBridge 暴露最小白名单 API。",
                            "不要把 ipcRenderer 或 Node fs/path/process/shell 直接暴露给页面。",
                            "主进程需要校验 IPC 消息来源和参数。",
                        ],
                    ),
                ],
                [
                    {"项目": "nodeIntegration", "建议": "false", "原因": "页面不应获得 Node 能力"},
                    {"项目": "contextIsolation", "建议": "true", "原因": "preload 与网页隔离"},
                    {"项目": "sandbox", "建议": "true", "原因": "降低渲染进程特权"},
                ],
            ),
        ),
        FileFixture(
            title="FastAPI TestClient 多格式摘要",
            filename="FastAPI TestClient 多格式摘要.pdf",
            media_type="application/pdf",
            expected_source_type="pdf_text",
            content=make_text_pdf(
                [
                    "FastAPI TestClient multiformat summary. Import path: fastapi.testclient TestClient. Wrapper: TestClient(app).",
                    "Usage: client.get and client.post call routes in tests. Boundary: no separate uvicorn process is required.",
                ]
            ),
        ),
        FileFixture(
            title="SQLite FTS5 BM25 多格式摘要",
            filename="SQLite FTS5 BM25 多格式摘要.md",
            media_type="text/markdown",
            expected_source_type="markdown",
            content=bytes_utf8(
                """
                # SQLite FTS5 BM25 多格式摘要

                ## 全文检索
                创建方式：可以用 CREATE VIRTUAL TABLE docs USING fts5(title, body) 创建 FTS5 虚表。
                查询方式：全文检索使用 MATCH 表达式，例如 docs MATCH 'sqlite'。
                排序函数：bm25(fts_table) 可以作为相关性排序信号。
                排序方向：FTS5 的 bm25 分数越小表示匹配越好，常见写法是 ORDER BY bm25(fts_table)。
                删除边界：删除资料后不能让旧 chunk 继续参与 MATCH。

                ## 示例表
                | 项目 | 写法 | 用途 |
                | --- | --- | --- |
                | 虚表 | CREATE VIRTUAL TABLE docs USING fts5(title, body) | 建立全文索引 |
                | 查询 | docs MATCH 'sqlite' | 过滤匹配文档 |
                | 排序 | ORDER BY bm25(docs) | 让更相关结果排在前面 |
                """
            ),
        ),
        FileFixture(
            title="Ollama Embeddings 多格式摘要",
            filename="Ollama Embeddings 多格式摘要.txt",
            media_type="text/plain",
            expected_source_type="plain_text_file",
            content=bytes_utf8(
                """
                Ollama Embeddings 多格式摘要

                本地地址：Ollama 默认在 http://127.0.0.1:11434 提供本机服务。
                端点：/api/embed 用于从输入文本生成 embedding 向量。
                请求字段：请求体包含 model 和 input。
                输入形态：input 可以是一段文本，也可以是多段文本列表。
                响应字段：响应里包含 embeddings 数组。
                安全边界：本地 Ollama embedding 不需要把资料片段发送给远程 embedding 服务商。
                """
            ),
        ),
        FileFixture(
            title="DeepSeek API 多格式摘要",
            filename="DeepSeek API 多格式摘要.pdf",
            media_type="application/pdf",
            expected_source_type="pdf_text",
            content=make_text_pdf(
                [
                    "DeepSeek API multiformat summary. Base URL: https://api.deepseek.com. Format: OpenAI compatible request format.",
                    "Chat endpoint: /chat/completions. Authentication: Authorization header with Bearer API Key. Logs must not record full keys.",
                    "RAG boundary: model requests may receive user question, chat context, and matched knowledge snippets.",
                ]
            ),
        ),
        FileFixture(
            title="MDN Fetch API 多格式摘要",
            filename="MDN Fetch API 多格式摘要.md",
            media_type="text/markdown",
            expected_source_type="markdown",
            content=bytes_utf8(
                """
                # MDN Fetch API 多格式摘要

                ## 请求与响应
                调用方式：fetch() 返回一个 Promise，成功时解析为 Response。
                HTTP 边界：404 或 500 这类 HTTP 错误状态不会自动让 Promise reject。
                错误判断：业务代码应检查 response.ok 或 response.status。
                读取方式：Response 的 json()、text() 等方法会读取响应体。
                请求体：发送 JSON 时通常设置 Content-Type: application/json，并把对象 JSON.stringify。
                """
            ),
        ),
        FileFixture(
            title="GitHub Actions Token 多格式摘要",
            filename="GitHub Actions Token 多格式摘要.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            expected_source_type="docx",
            content=make_docx_bytes(
                "GitHub Actions Token 多格式摘要",
                [
                    (
                        "权限控制",
                        [
                            "自动令牌：GitHub 会为 workflow job 提供 GITHUB_TOKEN。",
                            "权限声明：可以用 permissions 字段限制令牌权限。",
                            "最小权限：应该只给工作流需要的权限。",
                            "发布边界：上传 Release 资产通常需要 contents: write 权限。",
                            "安全边界：不要把令牌打印到日志，不要把令牌写进仓库文件。",
                        ],
                    )
                ],
                [
                    {"动作": "读取源码", "建议权限": "contents: read", "说明": "只读工作流使用"},
                    {"动作": "上传 Release 资产", "建议权限": "contents: write", "说明": "发布工作流使用"},
                ],
            ),
        ),
        FileFixture(
            title="Python venv 多格式摘要",
            filename="Python venv 多格式摘要.txt",
            media_type="text/plain",
            expected_source_type="plain_text_file",
            content=bytes_utf8(
                """
                Python venv 多格式摘要

                创建方式：python -m venv .venv 会创建一个独立虚拟环境目录。
                用途：虚拟环境拥有自己的 Python 可执行文件和 site-packages。
                激活方式：Windows PowerShell 常见激活脚本是 .venv\\Scripts\\Activate.ps1。
                依赖边界：项目依赖应安装进虚拟环境，避免污染系统 Python。
                删除边界：删除虚拟环境目录不会删除项目源码。
                """
            ),
        ),
    ]


PUBLIC_MULTIFORMAT_FIXTURES = build_public_multiformat_fixtures()


def build_multiformat_distractor_fixtures(count: int = 14) -> list[FileFixture]:
    topics = [
        ("浏览器存储", "localStorage 存储字符串键值，清理站点数据会影响保存内容。"),
        ("CSS Grid", "grid-template-columns 描述网格列轨道，gap 设置行列间距。"),
        ("pytest fixture", "fixture 可以为测试准备输入数据和清理逻辑。"),
        ("PowerShell 参数", "-NoProfile 可减少环境干扰，ExecutionPolicy 影响脚本运行策略。"),
        ("Pydantic 模型", "BaseModel 用字段类型定义输入输出结构。"),
        ("Node EventEmitter", "EventEmitter 使用 on 注册监听，emit 触发事件。"),
        ("Electron 打包", "本地候选包不应混入 SQLite、.env 或测试 fixture。"),
    ]
    fixtures: list[FileFixture] = []
    for index in range(1, count + 1):
        topic, description = topics[(index - 1) % len(topics)]
        title = f"多格式公开干扰资料 {index:02d} {topic}"
        extension = ".md" if index % 2 else ".txt"
        media_type = "text/markdown" if extension == ".md" else "text/plain"
        source_type = "markdown" if extension == ".md" else "plain_text_file"
        content = f"""
        # {title}

        ## 摘要
        编号：MF-PUBLIC-DISTRACTOR-{index:02d}
        说明：{description}
        边界：这份资料只用于扩大多格式公开资料基准规模，不应替代目标官方摘要。
        """.strip()
        fixtures.append(
            FileFixture(
                title=title,
                filename=f"{title}{extension}",
                media_type=media_type,
                expected_source_type=source_type,
                content=bytes_utf8(content),
            )
        )
    return fixtures


PUBLIC_MULTIFORMAT_DISTRACTORS = build_multiformat_distractor_fixtures()
PUBLIC_MULTIFORMAT_BENCHMARK_FIXTURES = [*PUBLIC_MULTIFORMAT_FIXTURES, *PUBLIC_MULTIFORMAT_DISTRACTORS]


PUBLIC_MULTIFORMAT_LEXICAL_CASES = [
    BenchmarkCase("mf react cleanup order", "React useEffect 依赖变化时 cleanup 和 setup 顺序是什么？", "React useEffect 多格式摘要", ("先用旧值运行 cleanup", "再用新值运行 setup"), ("VITE_",)),
    BenchmarkCase("mf react unmount", "React useEffect 组件卸载后会运行什么？", "React useEffect 多格式摘要", ("最后运行 cleanup",), ("TestClient",)),
    BenchmarkCase("mf react compare", "React useEffect 依赖比较使用什么？", "React useEffect 多格式摘要", ("Object.is",), ("response.ok",)),
    BenchmarkCase("mf react boundary", "React useEffect 没有外部系统同步时通常需要吗？", "React useEffect 多格式摘要", ("通常不需要 Effect",), ("GITHUB_TOKEN",)),
    BenchmarkCase("mf vite prefix", "Vite 哪些变量会暴露给客户端？", "Vite 环境变量多格式摘要", ("VITE_",), ("nodeIntegration",)),
    BenchmarkCase("mf vite access", "Vite 源码怎么读取环境变量？", "Vite 环境变量多格式摘要", ("import.meta.env",), ("TestClient(app)",)),
    BenchmarkCase("mf vite type", "Vite 环境变量值是什么类型？", "Vite 环境变量多格式摘要", ("字符串",), ("Object.is",)),
    BenchmarkCase("mf vite restart", "Vite 修改 .env 后需要做什么？", "Vite 环境变量多格式摘要", ("重启开发服务器",), ("uvicorn",)),
    BenchmarkCase("mf electron no node", "Electron 加载远程内容时 nodeIntegration 应该怎么设置？", "Electron 安全边界多格式摘要", ("不要启用 nodeIntegration",), ("VITE_",)),
    BenchmarkCase("mf electron isolation", "Electron contextIsolation 的作用是什么？", "Electron 安全边界多格式摘要", ("preload 与网页运行在不同上下文",), ("MATCH",)),
    BenchmarkCase("mf electron preload", "Electron preload 不能直接暴露什么？", "Electron 安全边界多格式摘要", ("ipcRenderer", "Node fs/path/process/shell"), ("TestClient",)),
    BenchmarkCase("mf electron table", "Electron 安全表里 sandbox 的建议值是什么？", "Electron 安全边界多格式摘要", ("sandbox", "true"), ("contents: write",)),
    BenchmarkCase("mf fastapi import", "FastAPI TestClient 从哪里导入？", "FastAPI TestClient 多格式摘要", ("fastapi.testclient", "TestClient"), ("contextBridge",)),
    BenchmarkCase("mf fastapi wrap", "FastAPI TestClient 怎么包装应用？", "FastAPI TestClient 多格式摘要", ("TestClient(app)",), ("CREATE VIRTUAL TABLE",)),
    BenchmarkCase("mf fastapi no uvicorn", "FastAPI TestClient 需要单独启动 uvicorn 吗？", "FastAPI TestClient 多格式摘要", ("no separate uvicorn process is required",), ("publish never",)),
    BenchmarkCase("mf fastapi methods", "FastAPI TestClient 用哪些方法调用路径？", "FastAPI TestClient 多格式摘要", ("client.get", "client.post"), ("response.ok",)),
    BenchmarkCase("mf sqlite create", "SQLite FTS5 怎么创建全文虚表？", "SQLite FTS5 BM25 多格式摘要", ("CREATE VIRTUAL TABLE docs USING fts5(title, body)",), ("TestClient",)),
    BenchmarkCase("mf sqlite match", "SQLite FTS5 查询用什么表达式？", "SQLite FTS5 BM25 多格式摘要", ("MATCH",), ("fetch()",)),
    BenchmarkCase("mf sqlite bm25", "SQLite FTS5 bm25 排序方向是什么？", "SQLite FTS5 BM25 多格式摘要", ("分数越小表示匹配越好", "ORDER BY bm25"), ("VITE_",)),
    BenchmarkCase("mf sqlite delete", "SQLite FTS5 删除资料后旧 chunk 还能参与 MATCH 吗？", "SQLite FTS5 BM25 多格式摘要", ("不能让旧 chunk 继续参与 MATCH",), ("cleanup",)),
    BenchmarkCase("mf ollama address", "Ollama 默认本地服务地址是什么？", "Ollama Embeddings 多格式摘要", ("http://127.0.0.1:11434",), ("https://api.deepseek.com",)),
    BenchmarkCase("mf ollama endpoint", "Ollama 生成 embedding 用哪个端点？", "Ollama Embeddings 多格式摘要", ("/api/embed",), ("/chat/completions",)),
    BenchmarkCase("mf ollama fields", "Ollama /api/embed 请求体包含什么？", "Ollama Embeddings 多格式摘要", ("model", "input"), ("GITHUB_TOKEN",)),
    BenchmarkCase("mf ollama privacy", "本地 Ollama embedding 隐私边界是什么？", "Ollama Embeddings 多格式摘要", ("不需要把资料片段发送给远程 embedding 服务商",), ("Bearer API Key",)),
    BenchmarkCase("mf deepseek base", "DeepSeek API 的 Base URL 是什么？", "DeepSeek API 多格式摘要", ("https://api.deepseek.com",), ("http://127.0.0.1:11434",)),
    BenchmarkCase("mf deepseek chat", "DeepSeek Chat Completions 使用哪个路径？", "DeepSeek API 多格式摘要", ("/chat/completions",), ("/api/embed",)),
    BenchmarkCase("mf deepseek auth", "DeepSeek API 认证 header 怎么带？", "DeepSeek API 多格式摘要", ("Authorization header", "Bearer API Key"), ("GITHUB_TOKEN",)),
    BenchmarkCase("mf deepseek rag", "DeepSeek RAG 请求可能收到什么？", "DeepSeek API 多格式摘要", ("user question", "chat context", "matched knowledge snippets"), ("Object.is",)),
    BenchmarkCase("mf fetch promise", "MDN Fetch API 中 fetch() 成功时返回什么？", "MDN Fetch API 多格式摘要", ("Promise", "Response"), ("TestClient(app)",)),
    BenchmarkCase("mf fetch http", "fetch 遇到 404 或 500 会自动 reject 吗？", "MDN Fetch API 多格式摘要", ("不会自动让 Promise reject",), ("cleanup",)),
    BenchmarkCase("mf fetch ok", "fetch 业务代码应该检查什么判断错误？", "MDN Fetch API 多格式摘要", ("response.ok", "response.status"), ("contents: write",)),
    BenchmarkCase("mf fetch json", "fetch 发送 JSON 通常怎么处理请求体？", "MDN Fetch API 多格式摘要", ("Content-Type: application/json", "JSON.stringify"), ("Object.is",)),
    BenchmarkCase("mf github token", "GitHub Actions job 会自动提供什么令牌？", "GitHub Actions Token 多格式摘要", ("GITHUB_TOKEN",), ("Bearer API Key",)),
    BenchmarkCase("mf github permissions", "GitHub Actions 怎么限制自动令牌权限？", "GitHub Actions Token 多格式摘要", ("permissions",), ("nodeIntegration",)),
    BenchmarkCase("mf github release", "GitHub Actions 上传 Release 资产通常需要什么权限？", "GitHub Actions Token 多格式摘要", ("contents: write",), ("Object.is",)),
    BenchmarkCase("mf github no log", "GitHub Actions 令牌能打印到日志吗？", "GitHub Actions Token 多格式摘要", ("不要把令牌打印到日志",), ("VITE_",)),
    BenchmarkCase("mf python venv create", "Python venv 怎么创建 .venv？", "Python venv 多格式摘要", ("python -m venv .venv",), ("TestClient",)),
    BenchmarkCase("mf python activate", "Windows PowerShell 激活 venv 的脚本是什么？", "Python venv 多格式摘要", (".venv\\Scripts\\Activate.ps1",), ("latest.yml",)),
    BenchmarkCase("mf python isolation", "Python venv 如何避免污染系统 Python？", "Python venv 多格式摘要", ("避免污染系统 Python",), ("VITE_",)),
    BenchmarkCase("mf generic base url", "Base URL 是什么？", None, should_inject=False, needs_clarification=True),
    BenchmarkCase("mf generic endpoint", "端点是什么？", None, should_inject=False, needs_clarification=True),
    BenchmarkCase("mf generic permission", "权限是什么？", None, should_inject=False, needs_clarification=True),
    BenchmarkCase("mf unrelated dinner", "今晚适合吃什么？", None, should_inject=False),
    BenchmarkCase("mf unrelated travel", "周末适合去哪旅游？", None, should_inject=False),
    BenchmarkCase("mf distractor css", "多格式公开干扰资料 02 CSS Grid 的编号是什么？", "多格式公开干扰资料 02 CSS Grid", ("MF-PUBLIC-DISTRACTOR-02",), ("MF-PUBLIC-DISTRACTOR-03",)),
    BenchmarkCase("mf distractor pydantic", "MF-PUBLIC-DISTRACTOR-05 Pydantic 模型说明什么？", "多格式公开干扰资料 05 Pydantic 模型", ("BaseModel",), ("EventEmitter",)),
    BenchmarkCase("mf distractor generic", "这些多格式公开干扰资料的编号是什么？", None, should_inject=False, needs_clarification=True),
]


PUBLIC_MULTIFORMAT_HYBRID_CASES = [
    BenchmarkCase("mf hybrid react cleanup", "React 副作用重新同步前会先跑清理函数吗？", "React useEffect 多格式摘要", ("先用旧值运行 cleanup",), ("VITE_",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid vite expose", "前端代码能直接看到哪些 Vite env？", "Vite 环境变量多格式摘要", ("VITE_",), ("GITHUB_TOKEN",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid electron bridge", "Electron 页面能不能直接拿 ipcRenderer 和 fs？", "Electron 安全边界多格式摘要", ("不要把 ipcRenderer 或 Node fs/path/process/shell 直接暴露给页面",), ("TestClient",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid fastapi process", "FastAPI 自动化测试是不是必须先起 uvicorn？", "FastAPI TestClient 多格式摘要", ("no separate uvicorn process is required",), ("publish never",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid sqlite ranking", "SQLite 全文检索哪个 bm25 分数更靠前？", "SQLite FTS5 BM25 多格式摘要", ("分数越小表示匹配越好",), ("Object.is",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid ollama local", "想在本机算向量不发到云端，Ollama 摘要怎么说？", "Ollama Embeddings 多格式摘要", ("不需要把资料片段发送给远程 embedding 服务商",), ("Bearer API Key",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid deepseek compatible", "DeepSeek API OpenAI compatible request format 使用哪个 Chat endpoint?", "DeepSeek API 多格式摘要", ("OpenAI compatible", "/chat/completions"), ("GITHUB_TOKEN",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid fetch status", "浏览器请求拿到 500 时是不是一定进 catch？", "MDN Fetch API 多格式摘要", ("不会自动让 Promise reject", "response.ok"), ("contents: write",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid github least privilege", "Actions 发布包时令牌权限怎么收紧？", "GitHub Actions Token 多格式摘要", ("permissions", "contents: write"), ("Bearer API Key",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid venv isolation", "Python 项目依赖怎么避免装进系统解释器？", "Python venv 多格式摘要", ("虚拟环境", "避免污染系统 Python"), ("latest.yml",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid docx table", "Electron 安全表里 contextIsolation 推荐值是什么？", "Electron 安全边界多格式摘要", ("contextIsolation", "true"), ("contents: write",), retrieval_mode="hybrid"),
    BenchmarkCase("mf hybrid github release table", "发布 Release 资产那行建议用什么权限？", "GitHub Actions Token 多格式摘要", ("上传 Release 资产", "contents: write"), ("Object.is",), retrieval_mode="hybrid"),
]


def seed_public_multiformat_files(client: TestClient) -> list[dict]:
    created: list[dict] = []
    for fixture in PUBLIC_MULTIFORMAT_BENCHMARK_FIXTURES:
        response = client.post(
            "/knowledge/import/file",
            files={"file": (fixture.filename, fixture.content, fixture.media_type)},
        )
        assert response.status_code == 201, f"{fixture.filename}: HTTP {response.status_code} {response.text}"
        data = response.json()
        assert data["title"] == fixture.title
        assert data["sourceType"] == fixture.expected_source_type
        assert data["chunkCount"] >= 1
        created.append(data)
    return created


def evaluate_public_multiformat_case(client: TestClient, case: BenchmarkCase, *, runtime: dict | None = None) -> str | None:
    return evaluate_case(client, case, runtime=runtime)


def test_public_multiformat_files_lexical_precision_and_no_answer_cases(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(tmp_path / "rag-public-multiformat.sqlite"))

    assert len(PUBLIC_MULTIFORMAT_FIXTURES) == 10
    assert len(PUBLIC_MULTIFORMAT_BENCHMARK_FIXTURES) >= 24
    assert len(PUBLIC_MULTIFORMAT_LEXICAL_CASES) >= 45

    with TestClient(app) as client:
        seed_public_multiformat_files(client)
        failures = [
            failure
            for case in PUBLIC_MULTIFORMAT_LEXICAL_CASES
            if (failure := evaluate_public_multiformat_case(client, case))
        ]

    assert not failures, "\n".join(failures)


def test_public_multiformat_files_hybrid_mock_embeddings(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(tmp_path / "rag-public-multiformat.sqlite"))
    runtime = public_multiformat_runtime_config()

    assert len(PUBLIC_MULTIFORMAT_HYBRID_CASES) >= 12

    with TestClient(app) as client:
        seed_public_multiformat_files(client)
        reindex = client.post("/knowledge/embeddings/reindex", json={"embeddingRuntimeConfig": runtime})
        assert reindex.status_code == 200
        assert reindex.json()["failed"] == 0
        assert reindex.json()["indexed"] >= len(PUBLIC_MULTIFORMAT_BENCHMARK_FIXTURES)

        failures = [
            failure
            for case in PUBLIC_MULTIFORMAT_HYBRID_CASES
            if (failure := evaluate_public_multiformat_case(client, case, runtime=runtime))
        ]

    assert not failures, "\n".join(failures)


def test_public_multiformat_file_metadata_and_table_rows(tmp_path: Path, monkeypatch) -> None:
    db_path = tmp_path / "rag-public-multiformat.sqlite"
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(db_path))

    with TestClient(app) as client:
        seed_public_multiformat_files(client)
        source_counts = {source["sourceType"]: 0 for source in client.get("/knowledge/sources").json()}
        for source in client.get("/knowledge/sources").json():
            source_counts[source["sourceType"]] = source_counts.get(source["sourceType"], 0) + 1

        electron_search = client.post("/knowledge/search", json={"query": "Electron 安全表里 contextIsolation 推荐值是什么？", "topK": 5})
        assert electron_search.status_code == 200
        electron_hit = next(
            hit
            for hit in electron_search.json()["hits"]
            if hit["sourceTitle"] == "Electron 安全边界多格式摘要"
            and hit["metadata"].get("tableIndex") == 1
            and hit["metadata"].get("rowIndex") == 2
        )
        electron_metadata = electron_hit["metadata"]
        assert electron_metadata["sourceFormat"] == "docx"
        assert electron_metadata["fileName"] == "Electron 安全边界多格式摘要.docx"
        assert electron_metadata["tableIndex"] == 1
        assert electron_metadata["rowIndex"] == 2

        fastapi_search = client.post("/knowledge/search", json={"query": "FastAPI TestClient 怎么包装应用？", "topK": 3})
        assert fastapi_search.status_code == 200
        fastapi_hit = next(hit for hit in fastapi_search.json()["hits"] if hit["sourceTitle"] == "FastAPI TestClient 多格式摘要")
        fastapi_metadata = fastapi_hit["metadata"]
        assert fastapi_metadata["sourceFormat"] == "pdf_text"
        assert fastapi_metadata["fileName"] == "FastAPI TestClient 多格式摘要.pdf"
        assert fastapi_metadata["page"] in {1, 2}

    assert source_counts["markdown"] >= 2
    assert source_counts["plain_text_file"] >= 2
    assert source_counts["pdf_text"] >= 2
    assert source_counts["docx"] >= 2

    with sqlite3.connect(db_path) as connection:
        connection.row_factory = sqlite3.Row
        electron_rows = connection.execute(
            """
            SELECT c.heading_path, c.chunk_type, c.content, c.metadata_json
            FROM knowledge_chunks c
            JOIN knowledge_sources s ON s.id = c.source_id
            WHERE s.title = 'Electron 安全边界多格式摘要'
            ORDER BY c.chunk_index
            """
        ).fetchall()
        fastapi_rows = connection.execute(
            """
            SELECT c.heading_path, c.chunk_type, c.content, c.metadata_json
            FROM knowledge_chunks c
            JOIN knowledge_sources s ON s.id = c.source_id
            WHERE s.title = 'FastAPI TestClient 多格式摘要'
            ORDER BY c.chunk_index
            """
        ).fetchall()

    assert electron_rows
    table_row = next(row for row in electron_rows if "contextIsolation" in row["content"] and "建议：true" in row["content"])
    table_metadata = json.loads(table_row["metadata_json"])
    assert table_row["chunk_type"] == "fact_block"
    assert table_metadata["sourceFormat"] == "docx"
    assert table_metadata["fileName"] == "Electron 安全边界多格式摘要.docx"
    assert table_metadata["tableIndex"] == 1
    assert PUBLIC_MULTIFORMAT_SOURCE_URLS["Electron 安全边界多格式摘要"].startswith("https://www.electronjs.org/")

    assert fastapi_rows
    pdf_row = next(row for row in fastapi_rows if "TestClient(app)" in row["content"])
    pdf_metadata = json.loads(pdf_row["metadata_json"])
    assert pdf_metadata["sourceFormat"] == "pdf_text"
    assert pdf_metadata["fileName"] == "FastAPI TestClient 多格式摘要.pdf"
    assert pdf_metadata["page"] in {1, 2}


def test_public_multiformat_deleted_file_is_not_recalled(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(tmp_path / "rag-public-multiformat.sqlite"))

    with TestClient(app) as client:
        created = seed_public_multiformat_files(client)
        target = next(source for source in created if source["title"] == "DeepSeek API 多格式摘要")
        before_delete = client.post("/knowledge/search", json={"query": "DeepSeek API 的 Base URL 是什么？", "topK": 3})
        assert before_delete.status_code == 200
        assert before_delete.json()["shouldInject"] is True

        deleted = client.delete(f"/knowledge/sources/{target['id']}")
        assert deleted.status_code == 200

        after_delete = client.post("/knowledge/search", json={"query": "DeepSeek API 的 Base URL 是什么？", "topK": 3})
        assert after_delete.status_code == 200
        data = after_delete.json()
        assert data["hits"] == []
        assert data["promptContext"] == ""
        assert data["shouldInject"] is False
