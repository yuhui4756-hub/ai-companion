from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from backend.app.document_parsing import DocumentParseError, parse_document_bytes
from backend.app.main import app
from backend.app.schemas import KnowledgeSourceType


def make_text_pdf(pages: list[str]) -> bytes:
    objects: list[bytes] = []
    page_ids: list[int] = []
    content_ids: list[int] = []

    for index in range(len(pages)):
        page_ids.append(3 + index * 2)
        content_ids.append(4 + index * 2)

    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(
        f"<< /Type /Pages /Kids [{' '.join(f'{page_id} 0 R' for page_id in page_ids)}] /Count {len(page_ids)} >>".encode(
            "ascii"
        )
    )

    for page_id, content_id, text in zip(page_ids, content_ids, pages):
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> /Contents {content_id} 0 R >>".encode(
                "ascii"
            )
        )
        text_ops = []
        for line in text.splitlines() or [""]:
            escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            text_ops.append(f"({escaped}) Tj 0 -16 Td")
        stream = f"BT /F1 12 Tf 72 720 Td {' '.join(text_ops)} ET".encode("ascii")
        objects.append(f"<< /Length {len(stream)} >>\nstream\n".encode("ascii") + stream + b"\nendstream")

    pdf = b"%PDF-1.4\n"
    offsets = [0]
    for obj_id, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += f"{obj_id} 0 obj\n".encode("ascii") + obj + b"\nendobj\n"
    xref_offset = len(pdf)
    pdf += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii")
    for offset in offsets[1:]:
        pdf += f"{offset:010d} 00000 n \n".encode("ascii")
    pdf += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode(
            "ascii"
        )
    )
    return pdf


def make_blank_pdf() -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def make_docx_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("月澜活动", level=1)
    document.add_heading("活动字段", level=2)
    document.add_paragraph("说明：月澜活动用于验证 DOCX 段落和表格解析。")
    document.add_paragraph("第一步：确认报名名单。", style="List Bullet")
    table = document.add_table(rows=1, cols=4)
    headers = ["编号", "预算金额", "负责人", "截止日期"]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    row = table.add_row().cells
    row[0].text = "YL-2026-009"
    row[1].text = "3.4 万元"
    row[2].text = "林澈"
    row[3].text = "2026-09-02"

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_parse_text_and_markdown_documents() -> None:
    text = parse_document_bytes("notes.txt", "普通文本资料\n编号：TXT-001".encode("utf-8"))
    markdown = parse_document_bytes("guide.md", "# 指南\n\n编号：MD-001".encode("utf-8"))

    assert text.source_type == KnowledgeSourceType.plain_text_file
    assert text.title == "notes"
    assert text.metadata["sourceFormat"] == "plain_text_file"
    assert "TXT-001" in text.content

    assert markdown.source_type == KnowledgeSourceType.markdown
    assert markdown.title == "guide"
    assert markdown.metadata["sourceFormat"] == "markdown"
    assert "# 指南" in markdown.content


def test_parse_document_rejects_unsupported_or_bad_text() -> None:
    with pytest.raises(DocumentParseError, match="仅支持"):
        parse_document_bytes("legacy.doc", b"not a docx")

    with pytest.raises(DocumentParseError, match="UTF-8"):
        parse_document_bytes("bad.txt", b"\xff\xfe\x00")


def test_parse_pdf_text_layer_and_reject_blank_pdf() -> None:
    parsed = parse_document_bytes(
        "starblue-plan.pdf",
        make_text_pdf(["PDFCASE-2026-041 Budget 18.4 wan", "Owner XuNian Deadline 2026-10-01"]),
    )

    assert parsed.source_type == KnowledgeSourceType.pdf_text
    assert parsed.metadata["sourceFormat"] == "pdf_text"
    assert parsed.metadata["pageCount"] == 2
    assert "## 第 1 页" in parsed.content
    assert "PDFCASE-2026-041" in parsed.content
    assert "## 第 2 页" in parsed.content

    with pytest.raises(DocumentParseError, match="OCR"):
        parse_document_bytes("blank.pdf", make_blank_pdf())


def test_parse_docx_paragraphs_lists_and_table_rows() -> None:
    parsed = parse_document_bytes("moon-event.docx", make_docx_bytes())

    assert parsed.source_type == KnowledgeSourceType.docx
    assert parsed.metadata["sourceFormat"] == "docx"
    assert "# 月澜活动" in parsed.content
    assert "## 活动字段" in parsed.content
    assert "- 第一步：确认报名名单。" in parsed.content
    assert "### 表格 1 第 1 行" in parsed.content
    assert "编号：YL-2026-009" in parsed.content
    assert "预算金额：3.4 万元" in parsed.content
    assert "负责人：林澈" in parsed.content


def test_file_import_endpoint_chunks_pdf_metadata_and_duplicate(tmp_path: Path, monkeypatch) -> None:
    db_path = tmp_path / "document-import.sqlite"
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(db_path))
    pdf_bytes = make_text_pdf(["PDFCASE-2026-041 Budget 18.4 wan Owner XuNian"])

    with TestClient(app) as client:
        created = client.post(
            "/knowledge/import/file",
            files={"file": ("starblue-plan.pdf", pdf_bytes, "application/pdf")},
        )
        assert created.status_code == 201
        data = created.json()
        assert data["title"] == "starblue-plan"
        assert data["sourceType"] == "pdf_text"
        assert data["chunkCount"] >= 1

        duplicate = client.post(
            "/knowledge/import/file",
            files={"file": ("starblue-plan.pdf", pdf_bytes, "application/pdf")},
        )
        assert duplicate.status_code == 409

        search = client.post("/knowledge/search", json={"query": "PDFCASE-2026-041", "topK": 3})
        assert search.status_code == 200
        assert search.json()["shouldInject"] is True
        assert "PDFCASE-2026-041" in search.json()["promptContext"]

    with sqlite3_connect(db_path) as connection:
        rows = connection.execute(
            """
            SELECT c.heading_path, c.chunk_type, c.content, c.metadata_json
            FROM knowledge_chunks c
            JOIN knowledge_sources s ON s.id = c.source_id
            WHERE s.title = 'starblue-plan'
            ORDER BY c.chunk_index
            """
        ).fetchall()

    assert rows
    target = next(row for row in rows if "PDFCASE-2026-041" in row["content"])
    metadata = json.loads(target["metadata_json"])
    assert metadata["sourceFormat"] == "pdf_text"
    assert metadata["fileName"] == "starblue-plan.pdf"
    assert metadata["page"] == 1
    assert metadata["contentHash"]


def test_file_import_endpoint_chunks_docx_table_row_metadata(tmp_path: Path, monkeypatch) -> None:
    db_path = tmp_path / "document-import.sqlite"
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(db_path))

    with TestClient(app) as client:
        created = client.post(
            "/knowledge/import/file",
            files={
                "file": (
                    "moon-event.docx",
                    make_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert created.status_code == 201
        assert created.json()["sourceType"] == "docx"

        search = client.post("/knowledge/search", json={"query": "moon-event YL-2026-009", "topK": 3})
        assert search.status_code == 200
        assert search.json()["shouldInject"] is True
        assert "YL-2026-009" in search.json()["promptContext"]
        assert "3.4 万元" in search.json()["promptContext"]

    with sqlite3_connect(db_path) as connection:
        rows = connection.execute(
            """
            SELECT c.heading_path, c.chunk_type, c.content, c.metadata_json
            FROM knowledge_chunks c
            JOIN knowledge_sources s ON s.id = c.source_id
            WHERE s.title = 'moon-event'
            ORDER BY c.chunk_index
            """
        ).fetchall()

    table_row = next(row for row in rows if "YL-2026-009" in row["content"])
    metadata = json.loads(table_row["metadata_json"])
    assert table_row["chunk_type"] == "fact_block"
    assert "预算金额：3.4 万元" in table_row["content"]
    assert metadata["sourceFormat"] == "docx"
    assert metadata["tableIndex"] == 1
    assert metadata["rowIndex"] == 1


def test_file_import_endpoint_rejects_large_and_unsupported_files(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(tmp_path / "document-import.sqlite"))

    with TestClient(app) as client:
        unsupported = client.post(
            "/knowledge/import/file",
            files={"file": ("image.png", b"fake image", "image/png")},
        )
        assert unsupported.status_code == 400
        assert "仅支持" in unsupported.json()["detail"]

        too_large = client.post(
            "/knowledge/import/file",
            files={"file": ("huge.txt", b"a" * (5 * 1024 * 1024 + 1), "text/plain")},
        )
        assert too_large.status_code == 413
        assert "文件过大" in too_large.json()["detail"]


def sqlite3_connect(path: Path):
    import sqlite3

    connection = sqlite3.connect(path)
    connection.row_factory = sqlite3.Row
    return connection
