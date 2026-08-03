from __future__ import annotations

import json
import sqlite3
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from fastapi.testclient import TestClient

from backend.app.main import app
from backend.tests.rag_case_loader import load_benchmark_cases
from backend.tests.test_document_parsing import make_text_pdf
from backend.tests.test_rag_realistic_benchmark import BenchmarkCase, evaluate_case


PUBLIC_M4C_ACCESSED_AT = "2026-08-03"
PUBLIC_M4C_CASE_FILE = Path(__file__).parent / "fixtures" / "rag_evidence_cases" / "public_m4c_cases.jsonl"


@dataclass(frozen=True)
class PublicM4CFileFixture:
    title: str
    filename: str
    content: bytes
    media_type: str
    expected_source_type: str
    source_url: str
    accessed_at: str = PUBLIC_M4C_ACCESSED_AT


def bytes_utf8(value: str) -> bytes:
    return value.strip().encode("utf-8")


def make_docx_bytes(title: str, sections: list[tuple[str, list[str]]], table_rows: list[dict[str, str]]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(title, level=1)
    for heading, lines in sections:
        document.add_heading(heading, level=2)
        for line in lines:
            document.add_paragraph(line)
    if table_rows:
        document.add_heading("表格字段", level=2)
        headers = list(table_rows[0].keys())
        table = document.add_table(rows=1, cols=len(headers))
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row_data in table_rows:
            row = table.add_row().cells
            for index, header in enumerate(headers):
                row[index].text = row_data[header]

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def public_m4c_runtime_config() -> dict:
    return {
        "providerName": "mock",
        "baseURL": "http://127.0.0.1:8765/mock",
        "model": "mock-embedding-public-m4c",
        "dimensions": 64,
        "batchSize": 4,
        "timeoutMs": 3000,
        "enabled": True,
        "apiKey": "mock-local-key",
    }


def build_public_m4c_core_fixtures() -> list[PublicM4CFileFixture]:
    return [
        PublicM4CFileFixture(
            title="Public M4C React useEffect 摘要",
            filename="Public M4C React useEffect 摘要.md",
            media_type="text/markdown",
            expected_source_type="markdown",
            source_url="https://react.dev/reference/react/useEffect",
            content=bytes_utf8(
                """
                # Public M4C React useEffect 摘要

                ## 副作用生命周期
                术语：useEffect 用来让组件和外部系统同步。
                setup 函数可以返回 cleanup 函数。
                依赖变化：React 会先用旧值运行 cleanup，再用新值运行 setup。
                依赖比较：React 使用 Object.is 比较 dependencies 中每一项。
                边界：没有外部系统需要同步时，通常不需要 Effect。
                """
            ),
        ),
        PublicM4CFileFixture(
            title="Public M4C React 不需要 Effect 摘要",
            filename="Public M4C React 不需要 Effect 摘要.txt",
            media_type="text/plain",
            expected_source_type="plain_text_file",
            source_url="https://react.dev/learn/you-might-not-need-an-effect",
            content=bytes_utf8(
                """
                Public M4C React 不需要 Effect 摘要

                派生值：如果一个值可以由 props 或 state 计算出来，不要用 Effect 和冗余 state 保存。
                昂贵计算：只有计算很昂贵时，才考虑用 useMemo 缓存计算结果。
                事件逻辑：由用户动作触发的逻辑应该放在 event handler，不放进 Effect。
                重置状态：需要重置整个子树 state 时，可以通过改变 key。
                """
            ),
        ),
        PublicM4CFileFixture(
            title="Public M4C MDN Fetch 摘要",
            filename="Public M4C MDN Fetch 摘要.md",
            media_type="text/markdown",
            expected_source_type="markdown",
            source_url="https://developer.mozilla.org/en-US/docs/Web/API/Fetch_API/Using_Fetch",
            content=bytes_utf8(
                """
                # Public M4C MDN Fetch 摘要

                ## 请求与响应
                调用方式：fetch() 返回一个 Promise，成功时解析为 Response。
                HTTP 边界：404 或 500 这类 HTTP 错误状态不会自动让 Promise reject。
                错误判断：业务代码应检查 response.ok 或 response.status。
                读取方式：Response 的 json()、text() 等方法会读取响应体。
                请求体：发送 JSON 时通常设置 Content-Type: application/json，并把对象 JSON.stringify。
                """
            ),
        ),
        PublicM4CFileFixture(
            title="Public M4C MDN Response.ok 摘要",
            filename="Public M4C MDN Response.ok 摘要.txt",
            media_type="text/plain",
            expected_source_type="plain_text_file",
            source_url="https://developer.mozilla.org/en-US/docs/Web/API/Response/ok",
            content=bytes_utf8(
                """
                Public M4C MDN Response.ok 摘要

                字段定义：Response.ok 是只读布尔属性。
                true 范围：状态码在 200 到 299 之间时，ok 为 true。
                使用建议：进入成功解析路径前，先检查 response.ok。
                边界：Response.ok 只说明 HTTP 状态是否在成功范围，不读取响应体。
                """
            ),
        ),
        PublicM4CFileFixture(
            title="Public M4C MDN Window.fetch 摘要",
            filename="Public M4C MDN Window.fetch 摘要.pdf",
            media_type="application/pdf",
            expected_source_type="pdf_text",
            source_url="https://developer.mozilla.org/en-US/docs/Web/API/Window/fetch",
            content=make_text_pdf(
                [
                    "Public M4C MDN Window.fetch summary. Window.fetch starts a request for a resource and returns a Promise.",
                    "Arguments: input is a RequestInfo or URL, and init is optional. Boundary: HTTP status details belong to Response.",
                ]
            ),
        ),
        PublicM4CFileFixture(
            title="Public M4C SQLite FTS5 摘要",
            filename="Public M4C SQLite FTS5 摘要.md",
            media_type="text/markdown",
            expected_source_type="markdown",
            source_url="https://www.sqlite.org/fts5.html",
            content=bytes_utf8(
                """
                # Public M4C SQLite FTS5 摘要

                ## 全文检索
                创建方式：可以用 CREATE VIRTUAL TABLE docs USING fts5(title, body) 创建 FTS5 虚表。
                查询方式：全文检索使用 MATCH 表达式，例如 docs MATCH 'sqlite'。
                排序函数：bm25(fts_table) 可以作为相关性排序信号。
                排序方向：FTS5 的 bm25 分数越小表示匹配越好，常见写法是 ORDER BY bm25(fts_table)。
                同步边界：external content table 需要和业务表同步，删除资料后旧 chunk 不得继续参与 MATCH。

                ## 示例表
                | 项目 | 写法 | 用途 |
                | --- | --- | --- |
                | 虚表 | CREATE VIRTUAL TABLE docs USING fts5(title, body) | 建立全文索引 |
                | 查询 | docs MATCH 'sqlite' | 过滤匹配文档 |
                | 排序 | ORDER BY bm25(docs) | 让更相关结果排在前面 |
                """
            ),
        ),
        PublicM4CFileFixture(
            title="Public M4C FastAPI UploadFile 摘要",
            filename="Public M4C FastAPI UploadFile 摘要.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            expected_source_type="docx",
            source_url="https://fastapi.tiangolo.com/tutorial/request-files/",
            content=make_docx_bytes(
                "Public M4C FastAPI UploadFile 摘要",
                [
                    (
                        "文件参数",
                        [
                            "声明方式：文件参数通常使用 UploadFile，并配合 File()。",
                            "读取方式：UploadFile 支持 async read() 读取文件内容。",
                            "内存边界：UploadFile 适合大文件，因为它使用 SpooledTemporaryFile。",
                        ],
                    ),
                ],
                [
                    {"项目": "filename", "建议": "客户端上传的原始文件名", "原因": "显示和记录来源"},
                    {"项目": "content_type", "建议": "上传文件的媒体类型", "原因": "判断解析路径"},
                    {"项目": "file", "建议": "底层文件对象", "原因": "流式处理内容"},
                ],
            ),
        ),
        PublicM4CFileFixture(
            title="Public M4C Electron contextIsolation 摘要",
            filename="Public M4C Electron contextIsolation 摘要.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            expected_source_type="docx",
            source_url="https://www.electronjs.org/docs/latest/tutorial/context-isolation",
            content=make_docx_bytes(
                "Public M4C Electron contextIsolation 摘要",
                [
                    (
                        "隔离边界",
                        [
                            "contextIsolation 会让 preload 脚本和网页脚本运行在不同上下文。",
                            "contextBridge 用来暴露经过挑选的最小 API。",
                            "安全边界：不要把完整 ipcRenderer 或 Node fs/path/process/shell 直接暴露给页面。",
                        ],
                    ),
                ],
                [
                    {"项目": "contextIsolation", "建议": "true", "原因": "preload 与网页上下文隔离"},
                    {"项目": "contextBridge", "建议": "只暴露最小 API", "原因": "减少页面可调用能力"},
                    {"项目": "ipcRenderer", "建议": "不要完整暴露", "原因": "避免扩大攻击面"},
                ],
            ),
        ),
        PublicM4CFileFixture(
            title="Public M4C Python sqlite3 摘要",
            filename="Public M4C Python sqlite3 摘要.txt",
            media_type="text/plain",
            expected_source_type="plain_text_file",
            source_url="https://docs.python.org/3/library/sqlite3.html",
            content=bytes_utf8(
                """
                Public M4C Python sqlite3 摘要

                连接方式：使用 sqlite3.connect(path) 打开数据库连接。
                参数绑定：SQL 查询应使用 ? 占位符传参，降低 SQL injection 风险。
                提交事务：完成写入后需要 commit() 保存事务。
                行访问：connection.row_factory = sqlite3.Row 后，可以按列名读取结果。
                边界：不要把用户密钥写入测试 fixture 或 SQLite 导出。
                """
            ),
        ),
    ]


def build_public_m4c_distractor_fixtures() -> list[PublicM4CFileFixture]:
    topics = [
        (
            "React useMemo 干扰摘要",
            "https://react.dev/reference/react/useMemo",
            "useMemo 可以缓存昂贵计算的结果，但不负责和外部系统同步，也不定义 cleanup 顺序。",
        ),
        (
            "React useState 干扰摘要",
            "https://react.dev/reference/react/useState",
            "useState 保存组件状态，set 函数会安排下一次渲染，不负责请求资源。",
        ),
        (
            "MDN localStorage 干扰摘要",
            "https://developer.mozilla.org/en-US/docs/Web/API/Window/localStorage",
            "localStorage 保存字符串键值，清理站点数据可能移除保存内容。",
        ),
        (
            "MDN URLSearchParams 干扰摘要",
            "https://developer.mozilla.org/en-US/docs/Web/API/URLSearchParams",
            "URLSearchParams 用于处理查询字符串，不判断 HTTP 状态范围。",
        ),
        (
            "SQLite transaction 干扰摘要",
            "https://www.sqlite.org/lang_transaction.html",
            "transaction 主题讨论 BEGIN、COMMIT、ROLLBACK 这样的事务控制语句。",
        ),
        (
            "BackgroundTasks 干扰摘要",
            "https://fastapi.tiangolo.com/tutorial/background-tasks/",
            "BackgroundTasks 在响应后执行任务，不代表文件参数，也不提供上传文件名字段。",
        ),
        (
            "Electron IPC 干扰摘要",
            "https://www.electronjs.org/docs/latest/tutorial/ipc",
            "ipcMain 和 ipcRenderer 用于进程通信，但这份摘要不建议把完整 ipcRenderer 暴露给页面。",
        ),
        (
            "Python pathlib 干扰摘要",
            "https://docs.python.org/3/library/pathlib.html",
            "pathlib.Path 用于路径处理，不提供 sqlite3.Row，也不执行 SQL 占位符绑定。",
        ),
        (
            "Vite env 干扰摘要",
            "https://vite.dev/guide/env-and-mode",
            "Vite 只把 VITE_ 前缀变量暴露给客户端代码，和 Response.ok 状态范围无关。",
        ),
        (
            "CSS Grid 干扰摘要",
            "https://developer.mozilla.org/en-US/docs/Web/CSS/CSS_grid_layout",
            "CSS Grid 使用 grid-template-columns 和 gap 布局，不处理 Fetch 响应。",
        ),
    ]
    fixtures: list[PublicM4CFileFixture] = []
    for index, (topic, source_url, description) in enumerate(topics, start=1):
        source_type = "markdown" if index % 2 else "plain_text_file"
        extension = ".md" if source_type == "markdown" else ".txt"
        title = f"Public M4C {topic}"
        fixtures.append(
            PublicM4CFileFixture(
                title=title,
                filename=f"{title}{extension}",
                media_type="text/markdown" if source_type == "markdown" else "text/plain",
                expected_source_type=source_type,
                source_url=source_url,
                content=bytes_utf8(
                    f"""
                    # {title}

                    ## 摘要
                    编号：PUBLIC-M4C-DISTRACTOR-{index:02d}
                    说明：{description}
                    边界：这份资料只作为 public-m4c 相似来源干扰，不应替代目标摘要。
                    """
                ),
            )
        )
    return fixtures


PUBLIC_M4C_CORE_FIXTURES = build_public_m4c_core_fixtures()
PUBLIC_M4C_DISTRACTOR_FIXTURES = build_public_m4c_distractor_fixtures()
PUBLIC_M4C_BENCHMARK_FIXTURES = [*PUBLIC_M4C_CORE_FIXTURES, *PUBLIC_M4C_DISTRACTOR_FIXTURES]


PUBLIC_M4C_LEXICAL_CASES = [
    BenchmarkCase("public-m4c-react-effect-cleanup-order", "React useEffect 依赖变化时 cleanup 和 setup 顺序是什么？", "Public M4C React useEffect 摘要", ("先用旧值运行 cleanup", "再用新值运行 setup"), ("冗余 state",)),
    BenchmarkCase("public-m4c-react-effect-object-is", "React useEffect 依赖比较使用什么？", "Public M4C React useEffect 摘要", ("Object.is",), ("response.ok",)),
    BenchmarkCase("public-m4c-react-effect-no-external", "Public M4C React useEffect 摘要里没有外部系统同步时通常需要 Effect 吗？", "Public M4C React useEffect 摘要", ("通常不需要 Effect",), ("event handler",)),
    BenchmarkCase("public-m4c-react-no-effect-derived-state", "React 不需要 Effect 摘要里派生值应该怎么处理？", "Public M4C React 不需要 Effect 摘要", ("不要用 Effect 和冗余 state 保存",), ("先用旧值运行 cleanup",)),
    BenchmarkCase("public-m4c-react-no-effect-event-handler", "由用户动作触发的逻辑应该放在哪里？", "Public M4C React 不需要 Effect 摘要", ("event handler",), ("MATCH",)),
    BenchmarkCase("public-m4c-fetch-http-error-no-reject", "fetch 遇到 404 或 500 会自动 reject 吗？", "Public M4C MDN Fetch 摘要", ("不会自动让 Promise reject",), ("Object.is",)),
    BenchmarkCase("public-m4c-fetch-body-readers", "Public M4C MDN Fetch 摘要里 Response 可以用哪些方法读取响应体？", "Public M4C MDN Fetch 摘要", ("json()", "text()"), ("sqlite3.Row",)),
    BenchmarkCase("public-m4c-response-ok-range", "Response.ok 在什么状态码范围为 true？", "Public M4C MDN Response.ok 摘要", ("200 到 299",), ("JSON.stringify",)),
    BenchmarkCase("public-m4c-response-ok-success-path", "进入成功解析路径前应先检查什么？", "Public M4C MDN Response.ok 摘要", ("response.ok",), ("SpooledTemporaryFile",)),
    BenchmarkCase("public-m4c-window-fetch-arguments", "Window.fetch 的 input 和 init 参数是什么？", "Public M4C MDN Window.fetch 摘要", ("RequestInfo", "init is optional"), ("200 到 299",)),
    BenchmarkCase("public-m4c-sqlite-fts-create", "SQLite FTS5 怎么创建 docs 全文虚表？", "Public M4C SQLite FTS5 摘要", ("CREATE VIRTUAL TABLE docs USING fts5(title, body)",), ("sqlite3.connect",)),
    BenchmarkCase("public-m4c-sqlite-fts-match", "SQLite FTS5 全文检索查询用什么表达式？", "Public M4C SQLite FTS5 摘要", ("MATCH",), ("UploadFile",)),
    BenchmarkCase("public-m4c-sqlite-bm25-direction", "SQLite FTS5 的 bm25 分数哪个更好？", "Public M4C SQLite FTS5 摘要", ("分数越小表示匹配越好",), ("transaction 控制提交和回滚",)),
    BenchmarkCase("public-m4c-sqlite-external-sync", "FTS5 external content table 的同步边界是什么？", "Public M4C SQLite FTS5 摘要", ("需要和业务表同步", "旧 chunk 不得继续参与 MATCH"), ("contextBridge",)),
    BenchmarkCase("public-m4c-fastapi-uploadfile-file", "Public M4C FastAPI UploadFile 摘要里文件参数通常用什么并配合什么？", "Public M4C FastAPI UploadFile 摘要", ("UploadFile", "File()"), ("BackgroundTasks",)),
    BenchmarkCase("public-m4c-fastapi-uploadfile-read", "UploadFile 怎么读取文件内容？", "Public M4C FastAPI UploadFile 摘要", ("async read()",), ("response.ok",)),
    BenchmarkCase("public-m4c-fastapi-uploadfile-content-type", "UploadFile 表格里 content_type 的建议和原因是什么？", "Public M4C FastAPI UploadFile 摘要", ("上传文件的媒体类型", "判断解析路径"), ("ipcRenderer",)),
    BenchmarkCase("public-m4c-electron-context-isolation", "Electron contextIsolation 会隔离什么？", "Public M4C Electron contextIsolation 摘要", ("preload 脚本和网页脚本运行在不同上下文",), ("sqlite3.Row",)),
    BenchmarkCase("public-m4c-electron-context-bridge", "Electron contextBridge 应该暴露什么？", "Public M4C Electron contextIsolation 摘要", ("经过挑选的最小 API",), ("MATCH",)),
    BenchmarkCase("public-m4c-electron-no-full-ipcrenderer", "Electron 页面不能直接暴露什么完整能力？", "Public M4C Electron contextIsolation 摘要", ("完整 ipcRenderer", "Node fs/path/process/shell"), ("JSON.stringify",)),
    BenchmarkCase("public-m4c-python-sqlite-placeholder", "Python sqlite3 查询应该用什么占位符降低注入风险？", "Public M4C Python sqlite3 摘要", ("? 占位符", "SQL injection"), ("useMemo",)),
    BenchmarkCase("public-m4c-python-sqlite-row-factory", "Python sqlite3 如何按列名读取结果？", "Public M4C Python sqlite3 摘要", ("connection.row_factory = sqlite3.Row",), ("Response.ok",)),
    BenchmarkCase("public-m4c-confuse-fetch-vs-response-ok", "Response.ok 的 true 范围是哪份资料说明的？", "Public M4C MDN Response.ok 摘要", ("Response.ok", "200 到 299"), ("不会自动让 Promise reject",)),
    BenchmarkCase("public-m4c-confuse-no-effect-vs-useeffect", "不要用冗余 state 保存派生值是哪份资料说明的？", "Public M4C React 不需要 Effect 摘要", ("冗余 state",), ("先用旧值运行 cleanup",)),
    BenchmarkCase("public-m4c-confuse-electron-bridge-vs-ipc", "Public M4C Electron contextIsolation 摘要里页面能力应该怎样最小暴露？", "Public M4C Electron contextIsolation 摘要", ("contextBridge", "最小 API"), ("ipcMain 和 ipcRenderer 用于进程通信",)),
    BenchmarkCase("public-m4c-confuse-fts-vs-transaction", "Public M4C SQLite FTS5 摘要里 bm25 排序方向是什么？", "Public M4C SQLite FTS5 摘要", ("FTS5", "分数越小表示匹配越好"), ("BEGIN、COMMIT、ROLLBACK",)),
    BenchmarkCase("public-m4c-confuse-uploadfile-vs-backgroundtasks", "Public M4C FastAPI UploadFile 摘要里 filename 字段有什么用途？", "Public M4C FastAPI UploadFile 摘要", ("客户端上传的原始文件名", "显示和记录来源"), ("响应后执行任务",)),
    BenchmarkCase("public-m4c-generic-status", "预算金额是多少？", None, should_inject=False, needs_clarification=True),
    BenchmarkCase("public-m4c-generic-parameter", "截止日期是什么？", None, should_inject=False, needs_clarification=True),
    BenchmarkCase("public-m4c-generic-return", "项目编号是什么？", None, should_inject=False, needs_clarification=True),
    BenchmarkCase("public-m4c-generic-owner", "负责人是谁？", None, should_inject=False, needs_clarification=True),
    BenchmarkCase("public-m4c-no-answer-weather", "明天上海会下雨吗？", None, should_inject=False),
    BenchmarkCase("public-m4c-no-answer-dinner", "今晚适合吃什么？", None, should_inject=False),
    BenchmarkCase("public-m4c-no-answer-movie", "帮我推荐一部电影可以吗？", None, should_inject=False),
]


PUBLIC_M4C_HYBRID_CASES = [
    BenchmarkCase("public-m4c-hybrid-effect-cleanup", "React 副作用重新同步前是不是先跑清理函数？", "Public M4C React useEffect 摘要", ("先用旧值运行 cleanup",), ("冗余 state",), retrieval_mode="hybrid"),
    BenchmarkCase("public-m4c-hybrid-derived-state", "能从 props 或 state 算出来的值还要再开 Effect 存一份吗？", "Public M4C React 不需要 Effect 摘要", ("不要用 Effect 和冗余 state 保存",), ("Object.is",), retrieval_mode="hybrid"),
    BenchmarkCase("public-m4c-hybrid-fetch-500", "浏览器请求拿到 500 时是不是一定进 catch？", "Public M4C MDN Fetch 摘要", ("不会自动让 Promise reject", "response.ok"), ("200 到 299",), retrieval_mode="hybrid"),
    BenchmarkCase("public-m4c-hybrid-ok-2xx", "HTTP 成功范围那个布尔字段什么时候为真？", "Public M4C MDN Response.ok 摘要", ("200 到 299", "ok 为 true"), ("Window.fetch starts",), retrieval_mode="hybrid"),
    BenchmarkCase("public-m4c-hybrid-fts-ranking", "SQLite 全文检索里 bm25 排名是不是数值越大越靠前？", "Public M4C SQLite FTS5 摘要", ("分数越小表示匹配越好",), ("BEGIN、COMMIT、ROLLBACK",), retrieval_mode="hybrid"),
    BenchmarkCase("public-m4c-hybrid-uploadfile-read", "FastAPI 上传文件想异步读内容，用 UploadFile 的哪个方法？", "Public M4C FastAPI UploadFile 摘要", ("async read()",), ("BackgroundTasks",), retrieval_mode="hybrid"),
    BenchmarkCase("public-m4c-hybrid-electron-bridge", "Electron 网页需要调用能力时，preload 应该怎么给最小接口？", "Public M4C Electron contextIsolation 摘要", ("contextBridge", "最小 API"), ("sqlite3.Row",), retrieval_mode="hybrid"),
    BenchmarkCase("public-m4c-hybrid-sqlite-parameterized", "Python 写 SQL 时怎么传参数才不容易被注入？", "Public M4C Python sqlite3 摘要", ("? 占位符", "SQL injection"), ("Response.ok",), retrieval_mode="hybrid"),
]


PUBLIC_M4C_ALL_CASES = [*PUBLIC_M4C_LEXICAL_CASES, *PUBLIC_M4C_HYBRID_CASES]


def seed_public_m4c_files(client: TestClient) -> list[dict]:
    created: list[dict] = []
    for fixture in PUBLIC_M4C_BENCHMARK_FIXTURES:
        response = client.post(
            "/knowledge/import/file",
            files={"file": (fixture.filename, fixture.content, fixture.media_type)},
        )
        assert response.status_code == 201, f"{fixture.filename}: HTTP {response.status_code} {response.text}"
        data = response.json()
        assert data["title"] == fixture.title
        assert data["sourceType"] == fixture.expected_source_type
        assert data["chunkCount"] >= 1
        created.append(data)
    return created


def evaluate_public_m4c_case(client: TestClient, case: BenchmarkCase, *, runtime: dict | None = None) -> str | None:
    return evaluate_case(client, case, runtime=runtime)


def test_public_m4c_case_file_is_runnable_and_matches_cases() -> None:
    loaded = load_benchmark_cases(PUBLIC_M4C_CASE_FILE, corpus_id="public-m4c")

    assert 30 <= len(loaded) <= 50
    assert [case.name for case in loaded] == [case.name for case in PUBLIC_M4C_ALL_CASES]
    assert any(case.retrieval_mode == "hybrid" for case in loaded)
    assert sum(1 for case in loaded if case.expected_source is None) >= 6


def test_public_m4c_files_lexical_precision_and_no_answer_cases(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(tmp_path / "rag-public-m4c.sqlite"))

    assert len(PUBLIC_M4C_CORE_FIXTURES) >= 8
    assert len(PUBLIC_M4C_DISTRACTOR_FIXTURES) >= 8
    assert len(PUBLIC_M4C_BENCHMARK_FIXTURES) >= 18
    assert len(PUBLIC_M4C_LEXICAL_CASES) >= 30

    with TestClient(app) as client:
        seed_public_m4c_files(client)
        failures = [failure for case in PUBLIC_M4C_LEXICAL_CASES if (failure := evaluate_public_m4c_case(client, case))]

    assert not failures, "\n".join(failures)


def test_public_m4c_files_hybrid_mock_embeddings(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(tmp_path / "rag-public-m4c.sqlite"))
    runtime = public_m4c_runtime_config()

    assert len(PUBLIC_M4C_HYBRID_CASES) >= 8

    with TestClient(app) as client:
        seed_public_m4c_files(client)
        reindex = client.post("/knowledge/embeddings/reindex", json={"embeddingRuntimeConfig": runtime})
        assert reindex.status_code == 200
        assert reindex.json()["failed"] == 0
        assert reindex.json()["indexed"] >= len(PUBLIC_M4C_BENCHMARK_FIXTURES)

        failures = [
            failure
            for case in PUBLIC_M4C_HYBRID_CASES
            if (failure := evaluate_public_m4c_case(client, case, runtime=runtime))
        ]

    assert not failures, "\n".join(failures)


def test_public_m4c_multiformat_metadata_and_source_refs(tmp_path: Path, monkeypatch) -> None:
    db_path = tmp_path / "rag-public-m4c.sqlite"
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(db_path))

    with TestClient(app) as client:
        seed_public_m4c_files(client)
        source_counts: dict[str, int] = {}
        for source in client.get("/knowledge/sources").json():
            source_counts[source["sourceType"]] = source_counts.get(source["sourceType"], 0) + 1

        upload_search = client.post("/knowledge/search", json={"query": "UploadFile 表格里 content_type 的建议和原因是什么？", "topK": 5})
        assert upload_search.status_code == 200
        upload_hit = next(
            hit
            for hit in upload_search.json()["hits"]
            if hit["sourceTitle"] == "Public M4C FastAPI UploadFile 摘要"
            and hit["metadata"].get("tableIndex") == 1
            and hit["metadata"].get("rowIndex") == 2
        )
        assert upload_hit["metadata"]["sourceFormat"] == "docx"
        assert upload_hit["metadata"]["fileName"] == "Public M4C FastAPI UploadFile 摘要.docx"

        window_fetch_search = client.post("/knowledge/search", json={"query": "Window.fetch 的 input 和 init 参数是什么？", "topK": 3})
        assert window_fetch_search.status_code == 200
        pdf_hit = next(
            hit
            for hit in window_fetch_search.json()["hits"]
            if hit["sourceTitle"] == "Public M4C MDN Window.fetch 摘要"
        )
        assert pdf_hit["metadata"]["sourceFormat"] == "pdf_text"
        assert pdf_hit["metadata"]["fileName"] == "Public M4C MDN Window.fetch 摘要.pdf"
        assert pdf_hit["metadata"]["page"] in {1, 2}

    assert source_counts["markdown"] >= 4
    assert source_counts["plain_text_file"] >= 4
    assert source_counts["pdf_text"] >= 1
    assert source_counts["docx"] >= 2
    assert all(fixture.source_url.startswith("https://") for fixture in PUBLIC_M4C_CORE_FIXTURES)
    assert all(fixture.accessed_at == PUBLIC_M4C_ACCESSED_AT for fixture in PUBLIC_M4C_BENCHMARK_FIXTURES)

    with sqlite3.connect(db_path) as connection:
        connection.row_factory = sqlite3.Row
        upload_rows = connection.execute(
            """
            SELECT c.heading_path, c.chunk_type, c.content, c.metadata_json
            FROM knowledge_chunks c
            JOIN knowledge_sources s ON s.id = c.source_id
            WHERE s.title = 'Public M4C FastAPI UploadFile 摘要'
            ORDER BY c.chunk_index
            """
        ).fetchall()
        fts_rows = connection.execute(
            """
            SELECT c.heading_path, c.chunk_type, c.content, c.metadata_json
            FROM knowledge_chunks c
            JOIN knowledge_sources s ON s.id = c.source_id
            WHERE s.title = 'Public M4C SQLite FTS5 摘要'
            ORDER BY c.chunk_index
            """
        ).fetchall()

    table_row = next(row for row in upload_rows if "content_type" in row["content"])
    table_metadata = json.loads(table_row["metadata_json"])
    assert table_row["chunk_type"] == "fact_block"
    assert table_metadata["sourceFormat"] == "docx"
    assert table_metadata["tableIndex"] == 1
    assert table_metadata["rowIndex"] == 2
    assert any(row["chunk_type"] == "table_row" and "ORDER BY bm25(docs)" in row["content"] for row in fts_rows)


def test_public_m4c_deleted_sources_are_not_recalled(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("SUOYI_BACKEND_DB_PATH", str(tmp_path / "rag-public-m4c.sqlite"))
    runtime = public_m4c_runtime_config()

    with TestClient(app) as client:
        created = seed_public_m4c_files(client)
        reindex = client.post("/knowledge/embeddings/reindex", json={"embeddingRuntimeConfig": runtime})
        assert reindex.status_code == 200
        assert reindex.json()["failed"] == 0

        for title, query in [
            ("Public M4C Python sqlite3 摘要", "Public M4C Python sqlite3 摘要 row_factory 列名访问"),
            ("Public M4C MDN Window.fetch 摘要", "Public M4C MDN Window.fetch 的 RequestInfo 参数是什么？"),
        ]:
            target = next(source for source in created if source["title"] == title)
            before_delete = client.post(
                "/knowledge/search",
                json={"query": query, "topK": 3, "retrievalMode": "hybrid", "embeddingRuntimeConfig": runtime},
            )
            assert before_delete.status_code == 200
            assert before_delete.json()["shouldInject"] is True

            deleted = client.delete(f"/knowledge/sources/{target['id']}")
            assert deleted.status_code == 200

            after_delete = client.post(
                "/knowledge/search",
                json={"query": query, "topK": 3, "retrievalMode": "hybrid", "embeddingRuntimeConfig": runtime},
            )
            assert after_delete.status_code == 200
            data = after_delete.json()
            assert data["hits"] == []
            assert data["promptContext"] == ""
            assert data["shouldInject"] is False
